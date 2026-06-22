#!/usr/bin/env python3
# 生成「镜面覆盖塑料件 技术资料中文翻译」docx —— 把三个产品的英文技术资料(PPT/docx)译成中文,便于报价。
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/Users/chensiyuan5/Desktop/Github/shenyuan_work/注塑项目/金榕/未报价 镜面覆盖塑料件/镜面覆盖塑料件_技术资料中文翻译.docx"

d = docx.Document()
d.styles["Normal"].font.name = "Microsoft YaHei"
d.styles["Normal"].font.size = Pt(10.5)


def H1(t):
    p = d.add_heading(t, level=1)
    return p

def H2(t):
    d.add_heading(t, level=2)

def P(t, bold=False, color=None):
    p = d.add_paragraph()
    r = p.add_run(t)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p

def bullet(t):
    d.add_paragraph(t, style="List Bullet")

def table(headers, rows):
    t = d.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    d.add_paragraph()


# ===== 封面 =====
title = d.add_heading("镜面覆盖塑料件 — 技术资料中文翻译（报价用）", level=0)
P("客户：上海金榕汽车零部件股份有限公司 / 终端 Mopar-FCA（Jeep、Ram）")
P("说明：本文件把英文技术资料(PPT tech sheet / docx)逐项译为中文，原始牌号/件号/规格保留英文。仅供报价参考，以客户原件为准。", color=(0x80,0x80,0x80))
P("含三个子产品：① F16 气坝 Spat ② KM 镜面保护罩 Mirror Protector ③ EJ 门袋绑带 Door Bags & Straps", bold=True)
d.add_paragraph()

# ============================================================
H1("一、F16 气坝 Spat（来源：F16 End Cap_Spat tech sheet_5-18-2026.pptx）")
P("项目：DT – F16 Ram Sports Truck 端盖(气坝/SPAT)开发 – Mopar（由金榕 JINRONG 发包）", bold=True)

H2("1. 产品定义与设计供货范围 — 零售底涂件（第2页）")
table(["项目", "内容（中文）"], [
    ["零件名称", "气坝套件 左+右（Spat Kit R&L）— primed 底涂，可直接喷漆"],
    ["新件号", "82219733AA"],
    ["车型 & 年款", "DT Sports Truck 带前保险杠（销售代码 MCE）；Stage 1(5.7L) & Stage 2A(6.4L) 基础版；MY 2027-2029"],
    ["应用", "外观件 & 功能件"],
    ["安装位置", "前保险杠左右两端；售后（经销商端）加装"],
    ["材料/颜色/纹理",
     "气坝 Spat：MS.50041 TPO（旧规格 MS-DC-256）Type C，颜色号 CPN 4981\n"
     "支架 Bracket Support：MS.50041 LBI's Hifax TYC 1168P，CPN 4998\n"
     "传感器支架 Sensor Bracket：MS.50041 LBI's Hifax TYC 1168P，CPN 4981"],
    ["年用量", "1,020（零售件 primed）"],
    ["其它", "Spat 为浅蓝色(light blue)；本页为零售服务件定义"],
])

H2("2. 产品定义 — 定制车间喷漆件（第3页）")
P("零件名称：气坝套件 左+右 — 已喷漆，车身同色。安装地点：墨西哥 Saltillo 的 UTM 定制车间（UTM Custom Shop）。材料/颜色/表面：同 82219733AA，但按下表喷成车身色。合计年用量 = 3,060。")
table(["件号 PN", "颜色描述", "年用量"], [
    ["68785495AA", "ZR6 - Molten Red Pearl Coat（熔岩红珠光漆）", "520"],
    ["68785506AA", "ZG9 - Serrano Green Metallic（塞拉诺绿金属漆）", "367"],
    ["68785507AA", "KXJ - Diamond Black Crystal（钻石黑水晶）", "704"],
    ["68785508AA", "PDN - Ceramic Grey Clear Coat（陶瓷灰清漆）", "581"],
    ["68785509AA", "GW7 - Bright White Clear Coat（亮白清漆）", "643"],
    ["68785510AA", "EYB - Detonator Yellow（引爆黄）", "245"],
])

H2("3. 工作范围（第4页）")
bullet("范围 = 零件的设计、开发、制造、PPAP 与供货。")
bullet("交付预期：零售件 82219733AA（primed 底涂）单件包装发往 Mopar 仓库（Depot）；定制车间件 = 6 个件号、车身色，发往墨西哥 Saltillo 的 Utilimaster 定制车间。")
bullet("尽量沿用现有相似量产件的连接方式与现有五金件。金榕将提供“direct buy（定向采购）letter”，供向现有量产供应商采购现有件。")
bullet("间隙与表面质量目标 = 与本车现有相似量产件一致。")
bullet("A 面 CAD 数据在定点后提供；供应商负责 B 面与连接点的开发，并负责新件 CAD 与 2D 图纸的创建。")
bullet("报价前如有疑问可申请技术评审(tech review)。")
bullet("零件级测试由供应商负责；系统级与整车级测试由金榕或供应商进行。性能/测试要求见随 SP 附的 DVP&R。")
bullet("供应商需供应市场宣传与测试用件（备 3 套）。定向采购件(directed buy)不需重新测试；尽量用已验证、有测试数据的产品以减少新测试。")

H2("4. 零件小结 — 其它要求（第5页）")
bullet("最终数据须提交 Mopar 评审/批准后方可启动(kick-off)。")
bullet("供应商承担设计责任，须提供该件 CAD 模型。")
bullet("测试费须按 DVP 逐项列出，另含工程设计开发(ED&D)费用。")
bullet("GD&T(几何尺寸与公差)文档与报价假设须遵循 PS-9611。")
bullet("供应商负责编写安装说明书(installation instructions)与 I-Sheet；套件含金榕提供的二维码卡（链接到电子版 I-sheet）。")
P("验证(VALIDATION)：须满足 DVP&R 及金榕标准。若某 DVP 项拟用替代数据(surrogate data)，报价时须说明、与工程评审，且数据不超过 3 年，除非 Mopar 产品开发出具偏差许可。")
P("样件与量产(Sample & SoP)：工程装配试配 3 件(primed)；按 DVP 做测试/验证；供应商质量 3 套；市场 1 套；PER 样件 2 种车身色。")

H2("5. 时间计划要求（第6页）")
bullet("供应商须提供详细时间表，覆盖 PO、设计、测试、模具、DV 测试、系统级 PV 测试、包装、PPAP、物流、发运至 Mopar 仓库等全部活动，标明每项责任人/组与依赖关系。")
bullet("时间表须考虑节假日、停产、海运、二次开发等余量，以满足上面整车时间要求；并附假设清单、风险点及对策。")
bullet("须附详细流程：原材料、模具设计/制造、零件制造/开发/包装的地点与责任来源、发运地。")

H2("6. 结构/尺寸（第7–11页，图纸为主）")
bullet("第7页：左侧件示意 — 标出 Spat（气坝）与 Bracket（支架）。")
bullet("第9页：左侧件六视图；主要外形尺寸约 425 / 400 / 410（mm）。")
bullet("第10页：垫片 Spacer — 内径 4.2mm、外径 15~25mm、厚 5mm；封边(Close Out) 尺寸约 370/350；附量产挡泥板(splash shield)仅供参考。")

H2("7. 紧固件清单（第11页）")
table(["件号 PN", "描述（中文）", "材料", "数量"], [
    ["06509800AA", "M-4.8 U型卡扣，端盖固定到轮罩衬板", "PS-11036", "4"],
])

# ============================================================
H1("二、KM 镜面保护罩 Mirror Protector（来源：KM tech-sheet.pptx 图片页 + Mirror protector-FINAL.docx）")
P("项目：Jeep – Mexico – Mirror Protectors，车型 KM", bold=True)

H2("1. 产品定义（PPT 第2页，图片）")
table(["项目", "内容（中文）"], [
    ["功能 Function", "防止玻璃后视镜被盗"],
    ["材料 Material（原文不译）", "ASA BLACK / ABS"],
    ["纹理 Texture", "COD.003（⚠ PPT为COD.003，FINAL.docx写COD.001，需向客户确认）"],
    ["胶水 Adhesive（原文不译）", "Epoxy（EP-20 / H-97M，见 DATA SHEET Adhesive.pdf）"],
    ["装配 Fitment", "与后视镜各调节位置至少留 2mm 间隙；须沿所有边缘贴合；不得干涉后视镜任何功能"],
    ["车型 Vehicle", "KM（FINAL.docx 写 B10 & C10/C16）"],
    ["颜色/品牌", "黑色 Black"],
    ["件号 P/N", "82219786AA"],
    ["重量 Weight", "PPT：TBD；FINAL.docx：0.21 kg/套"],
    ["每套件数", "2 件 — 左(LH) + 右(RH)"],
    ["工艺 Process", "注塑成型 Injection Molding"],
])

H2("2. 工作范围与工程开发（Statement of Work and ED.docx）")
bullet("选定供应商负责整套件的设计、开发、PPAP 与供货。")
bullet("须基于客户产品设计室提供的 Artwork(外观数据)开发；供应商做样件并装到客户提供的车上做整车认可，可能需多次迭代（复杂曲面常需微调）。")
bullet("供应商还负责服务件的 3D 数据、2D 图纸、I-sheet 的创建。")
bullet("薄膜/材料(film material)须满足客户规范；如无现成数据，供应商负责提供替代数据/样板或重新测试。")

# ============================================================
H1("三、EJ 门袋绑带 Door Bags & Straps（来源：27EJ Door Bags tech sheet.pptx + EJ Door bags DVP.xlsx）")

H2("1. 套件物料清单 KIT BOM（PPT 第7页）")
table(["序号", "零件名称（原文不译，采购用）", "数量", "包装盒尺寸(mm)"], [
    ["1", "Buckle", "4", "50*15*40"],
    ["2", "Nylon Cords/Straps", "2", "410*87*37"],
    ["3", "Fastener", "4", "M8*25"],
    ["4", "I-sheet", "1", "NA"],
])
P("附注：PPT 中“现有带扣与绑带到车门的连接方式”仅作参考图；图中尺寸(mm)为近似值，仅供参考。")

H2("2. 材料规格（EJ Door bags DVP.xlsx — 材料认证项）")
table(["子件（原文不译）", "材料规格(原文不译)", "说明"], [
    ["Buckle", "MS.50042 / 颜色 MS-JP-1-3", "FCA 材料标准号；深远拟用 ASA 注塑"],
    ["Nylon Cords/Straps", "MS-JZ-12-20", "长 410mm，采购件"],
    ["Fastener", "M8*25", "金属标准件；表面处理按 52605_07，等级/镀层资料客户未给"],
])
P("其它 DVP 关键测试：材料/颜色耐久(MS-JP-1-3)、耐热老化(85℃ 7天)、绑带拉脱强度(PF.90058)、回收要求(CS-9003)等 —— 实验费不含在报价内，但产品须满足。")

# ============================================================
H1("四、报价提示（综合）")
bullet("三个产品都是“总成/套件”交付：塑料件我方注塑，标准化件(绑带/紧固件/胶/支架)采购计入，整套一价。")
bullet("F16 Spat 的两个 Hifax 支架走 direct buy 定向采购(客户给采购letter)，不需我方开模；喷漆在墨西哥定制车间完成，零售件只供 primed 底涂。")
bullet("MS.5004x / MS-JZ / MS-JP 等均为 FCA/Mopar 材料标准号，不是某厂牌号；具体厂家/单价需询价确认。")
bullet("⚠ 待客户确认：KM 纹理 COD.001还是COD.003、KM 重量、各件颜色与年用量口径、Hifax 支架是否确为定向采购、紧固件等级镀层。")

d.save(OUT)
print("已保存:", OUT)
