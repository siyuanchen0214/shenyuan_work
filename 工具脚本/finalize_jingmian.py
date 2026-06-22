#!/usr/bin/env python3
# 1) 生成胶水数据表中文对照docx; 2) 删除已转双语的旧.xls; 3) 所有文件名/子文件夹名改中文。
import os, docx
from docx.shared import Pt, RGBColor

BASE = "/Users/chensiyuan5/Desktop/Github/shenyuan_work/注塑项目/金榕/未报价 镜面覆盖塑料件 2"
KM = BASE + "/｣ｨKM -Mirror Protector -2026｣ｩ---Technical Solution"
EJ = BASE + "/EJ Door Bags and Straps---Technical Solution"

# ---------- 1) 胶水数据表 中文对照 ----------
d = docx.Document()
d.styles["Normal"].font.name = "Microsoft YaHei"
def H(t): d.add_heading(t, level=1)
def P(t, zh=False):
    p=d.add_paragraph(); r=p.add_run(t)
    if zh: r.font.color.rgb=RGBColor(0x1F,0x6F,0xC0)
d.add_heading("胶水数据表 EP-20/H-97M —— 中文对照(原文见同名PDF)", 0)
P("DATA SHEET EP-20/H-97M"); P("【中】环氧结构胶 EP-20/H-97M 数据表", True)
H("描述 DESCRIPTION")
P("Fast gel time epoxy resin, 1:1 mix ratio by volume; semi-rigid, good impact resistance; good adhesion to metals, plastics, wood, concrete, asphalt, etc.")
P("【中】快速凝胶环氧树脂胶,按体积 1:1 混合;半刚性,抗冲击好;对金属、塑料、木材、混凝土、沥青等附着力好。", True)
H("规格 SPECIFICATIONS")
P("EP-20: 黑色 Black;H-97M: 半透明 Translucent。体积混合比 100:100;重量混合比 100:100。"
  "比重@25°C: EP-20=1.16, H-97M=1.15。粘度 cps: EP-20=11,000–14,000, H-97M=10,000–16,000;"
  "混合后粘度 10,000–15,000 cps。凝胶时间 Gel time: 18–20 分钟。", True)
P("注:测试在图示温度、相对湿度 50%-80%(可 40%-85%)下进行,湿度不影响固化与性能。", True)
H("物理性能 PHYSICAL PROPERTIES")
P("压缩强度 Compression(ASTM-D-695): 12,500 psi;拉伸力 Tension(ASTM-D-638): 8,000 psi;"
  "伸长率 Elongation(ASTM-D-638): 7%;硬度 Hardness(ASTM-D-2240): 65–70 D。", True)
H("使用说明 INSTRUCTIONS FOR USE")
P("1. 清洁基材,确保无尘、无油脂及任何影响粘接的污染物。\n"
  "2. 按体积或重量混合比量取两组分,最好用配胶设备混合。\n"
  "3. 将混合胶涂到待粘接/修补的零件上。\n"
  "4. 不要让胶长时间处于 15°C(59°F)以下。\n"
  "5. 使用前数分钟,不要让胶受阳光直射或 30°C(86°F)以上热源。", True)
H("固化 CURING")
P("室温 16–24 小时初固化;25°C 下至少 7–10 天完全固化。", True)
H("储存 STORAGE")
P("阴凉干燥处,15°C–40°C;未开封最长可存 24 个月。", True)
H("安全 SAFETY")
P("避免吸入雾气与蒸气(可能损害呼吸道);操作建议戴手套、护目镜与口罩;在通风良好处作业;"
  "入眼用大量清水冲洗15分钟;避免接触皮肤,接触后用肥皂水清洗再用清水冲。详见安全数据表(SDS)。", True)
P("厂家:A ADVANCE SOLUTIONS, S.A DE C.V.(墨西哥)Tel +52 55 5699-7394", True)
P("免责:数据基于厂家实验室与经验,实际使用条件不可控,厂家不对操作与应用负责,请自行验证。", True)
d.save(KM + "/胶水数据表EP-20_H-97M_中文对照.docx")
print("胶水中文对照 OK")

# ---------- 2) 删除旧 .xls(已转双语 .xlsx) ----------
for x in ["ADP RR-Mopar Interior .xls", "ATAR (2).xls", "EDD template.xls"]:
    p = EJ + "/" + x
    if os.path.exists(p): os.remove(p); print("删除旧xls:", x)

# ---------- 3) 文件改中文名 ----------
ren = {
  BASE + "/F16 End Cap_Spat  tech sheet_5-18-2026 .pptx": BASE + "/F16气坝Spat_技术资料(中英).pptx",
  KM + "/KM -Mirror Protector -2026-tech-sheet.pptx": KM + "/KM镜面保护罩_技术资料(中英).pptx",
  KM + "/Mirror protector - FINAL.docx": KM + "/镜面保护罩_技术说明(中英).docx",
  KM + "/Statement of Work and ED.docx": KM + "/工作说明书与工程开发SoW(中英).docx",
  KM + "/Mirror protector - DVP.xlsx": KM + "/镜面保护罩_设计验证计划DVP(中英).xlsx",
  KM + "/PIS Format.xlsm": KM + "/零件检验标准PIS(中英).xlsm",
  KM + "/DATA SHEET Adhesive.pdf": KM + "/胶水数据表EP-20_H-97M(原文).pdf",
  KM + "/SP drawings - 2.pdf": KM + "/SP图纸(原文·纯图无文字).pdf",
  EJ + "/27EJ Door Bags-tech sheet.pptx": EJ + "/EJ门袋绑带_技术资料(中英).pptx",
  EJ + "/EJ Door bags DVP.xlsx": EJ + "/EJ门袋绑带_设计验证计划DVP(中英).xlsx",
  EJ + "/ADP RR-Mopar Interior .xlsx": EJ + "/职责与权责表RR-ADP(中英).xlsx",
  EJ + "/ATAR (2).xlsx": EJ + "/外观技术批准报告ATAR(中英).xlsx",
  EJ + "/EDD template.xlsx": EJ + "/工程数据需求EDD模板(中英).xlsx",
  EJ + "/PIS Format (1).xlsm": EJ + "/零件检验标准PIS(中英).xlsm",
  EJ + "/27J Door Bag Straps-Box dimensions.pdf": EJ + "/门袋绑带_包装箱尺寸(原文).pdf",
}
for a, b in ren.items():
    if os.path.exists(a): os.rename(a, b); print("改名:", os.path.basename(b))

# ---------- 子文件夹改中文名(最后做) ----------
fold = {KM: BASE + "/KM镜面保护罩_技术资料", EJ: BASE + "/EJ门袋绑带_技术资料"}
for a, b in fold.items():
    if os.path.exists(a): os.rename(a, b); print("文件夹改名:", os.path.basename(b))
print("DONE")
