"""
用法：
    python3 核对乙方回传合同.py <乙方回传的合同.docx>

会自动根据文件名匹配原版快照，逐段比对，输出所有改动。
"""
import sys, json, hashlib, difflib
from docx import Document

SNAPSHOT = '/Users/chensiyuan5/Desktop/Github/shenyuan_work/注塑项目/金榕/合同原版快照_甲方发出版_20260701.json'

def sha256_of(path):
    with open(path, 'rb') as f:
        return hashlib.sha256(f.read()).hexdigest()

def extract_paragraphs(path):
    doc = Document(path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

def extract_tables(path):
    doc = Document(path)
    result = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                rows.append(cells)
        result.append(rows)
    return result

def main():
    if len(sys.argv) < 2:
        print('用法: python3 核对乙方回传合同.py <回传合同.docx>')
        sys.exit(1)

    supplier_path = sys.argv[1]

    with open(SNAPSHOT, encoding='utf-8') as f:
        snapshot = json.load(f)

    # 根据文件名模糊匹配原版
    fname = supplier_path.lower()
    matched_key = None
    for key in snapshot['contracts']:
        if 'gea' in fname or '顶棚' in fname:
            if 'GEA' in key:
                matched_key = key
                break
        if 'lucid' in fname or '地板' in fname or '舱盖' in fname:
            if 'Lucid' in key:
                matched_key = key
                break

    if not matched_key:
        # 如果只有一个合同就直接用
        keys = list(snapshot['contracts'].keys())
        if len(keys) == 1:
            matched_key = keys[0]
        else:
            print('无法自动匹配合同，请检查文件名是否包含 GEA/顶棚 或 Lucid/地板/舱盖')
            sys.exit(1)

    orig = snapshot['contracts'][matched_key]
    print(f'\n【原版】{matched_key}')
    print(f'  快照生成时间: {snapshot["generated_at"]}')
    print(f'  原版 SHA256: {orig["sha256"]}')

    # 1. 哈希比对
    new_hash = sha256_of(supplier_path)
    print(f'  回传 SHA256: {new_hash}')
    if new_hash == orig['sha256']:
        print('\n✅ 文件与原版完全一致，未发现任何改动。')
        return

    print('\n⚠️  SHA256 不同，正在逐段比对……\n')

    # 2. 段落比对
    orig_paras = orig['paragraphs']
    new_paras = extract_paragraphs(supplier_path)

    diff = list(difflib.unified_diff(orig_paras, new_paras, lineterm='', n=0))
    if not diff:
        print('正文段落：无差异')
    else:
        print('=== 正文段落差异 ===')
        for line in diff:
            if line.startswith('---') or line.startswith('+++') or line.startswith('@@'):
                continue
            if line.startswith('-'):
                print(f'  【原版】{line[1:]}')
            elif line.startswith('+'):
                print(f'  【回传】{line[1:]}')

    # 3. 表格比对
    orig_tables = orig['tables']
    new_tables = extract_tables(supplier_path)

    table_diffs = []
    for t_idx, (orig_t, new_t) in enumerate(zip(orig_tables, new_tables)):
        orig_flat = [str(r) for r in orig_t]
        new_flat  = [str(r) for r in new_t]
        tdiff = list(difflib.unified_diff(orig_flat, new_flat, lineterm='', n=0))
        if tdiff:
            table_diffs.append((t_idx, tdiff))

    if not table_diffs:
        print('表格内容：无差异')
    else:
        print('\n=== 表格差异 ===')
        for t_idx, tdiff in table_diffs:
            print(f'  表格{t_idx+1}:')
            for line in tdiff:
                if line.startswith('---') or line.startswith('+++') or line.startswith('@@'):
                    continue
                if line.startswith('-'):
                    print(f'    【原版】{line[1:]}')
                elif line.startswith('+'):
                    print(f'    【回传】{line[1:]}')

if __name__ == '__main__':
    main()
